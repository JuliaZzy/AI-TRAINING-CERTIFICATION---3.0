#!/usr/bin/env python3
"""从操作题.md + QB ipynb + 素材 docx 生成小程序操作题数据（评分与题目/填空对齐）。"""

import html as html_module
import json
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))
from doc_templates import DOC_TEMPLATES
from index_answers import get_index_answers
OPS_MD = ROOT / "qustions" / "理论题" / "操作题.md"
QUESTIONS_HTML_DIR = ROOT / "qustions"
QB_DIR = ROOT / "QB"
MATERIAL_DIR = ROOT / "素材"
OUT_DIR = ROOT / "miniprogram" / "data"

CHAPTER_NAMES = {
    1: "业务分析",
    2: "智能训练",
    3: "智能系统设计",
    4: "培训与指导",
}

BLANK_PATTERN = re.compile(r"_{3,}")
BULLET_CHARS = "\uf06c\uf0b7\uf0d8\uf0a7\uf0fc"
SQUARE_CHARS = "□■▪"
BULLET_DOT = "• "


def normalize_bullets(text: str) -> str:
    if not text or not isinstance(text, str):
        return text
    result = text
    for ch in BULLET_CHARS + SQUARE_CHARS:
        result = result.replace(ch + "\t", BULLET_DOT)
        result = re.sub(rf"(^|\n){re.escape(ch)}\s*", r"\1" + BULLET_DOT, result)
        result = result.replace(ch, BULLET_DOT)
    while "• • " in result:
        result = result.replace("• • ", BULLET_DOT)
    return result


def normalize_deep(obj):
    if isinstance(obj, str):
        return normalize_bullets(obj)
    if isinstance(obj, list):
        return [normalize_deep(item) for item in obj]
    if isinstance(obj, dict):
        return {key: normalize_deep(value) for key, value in obj.items()}
    return obj


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def extract_field(block: str, label: str) -> str:
    pattern = rf"{re.escape(label)}\n(.*?)(?=\n\d+\.|$)"
    match = re.search(pattern, block, re.DOTALL)
    return match.group(1).strip() if match else ""


def parse_rubric(block: str) -> list[dict]:
    items = []
    for line in block.splitlines():
        line = line.strip()
        if not line.startswith("M") and not line.startswith("\t"):
            continue
        if line.startswith("M"):
            parts = line.split("\t")
            if len(parts) >= 3:
                score_match = re.search(r"(\d+)", parts[1])
                items.append({
                    "id": parts[0].strip(),
                    "score": int(score_match.group(1)) if score_match else 0,
                    "desc": parts[2].strip(),
                })
        elif line.startswith("\t") and "内容点" in line:
            score_match = re.search(r"(\d+)分", line)
            if score_match and items:
                items.append({
                    "id": items[-1]["id"] + "_sub",
                    "score": int(score_match.group(1)),
                    "desc": line.strip(),
                    "parent": items[-1]["id"],
                })
    total = 0
    for row in items:
        if "parent" not in row:
            total += row["score"]
    return items


def split_operation_md(text: str) -> list[tuple[str, str]]:
    sections = re.split(r"(?=人工智能训练师（三级）操作技能考核)", text)
    sheets: dict[str, str] = {}
    rubrics: dict[str, str] = {}

    for sec in sections:
        if not sec.strip():
            continue
        name_match = re.search(r"试题名称[:：](.+)", sec)
        if not name_match:
            continue
        name = name_match.group(1).strip()
        if "测量分评分表" in sec:
            rubrics[name] = sec
        elif "1.场地设备要求" in sec or "2.工作任务" in sec:
            sheets[name] = sec

    pairs = []
    for name, sheet in sheets.items():
        pairs.append((name, sheet, rubrics.get(name, "")))
    return pairs


def infer_question_id(name: str, tasks: str, rubric: list[dict]) -> str:
    for text in [tasks] + [r["desc"] for r in rubric] + [name]:
        for pattern in [
            r"(\d+\.\d+\.\d+)\.ipynb",
            r"素材(\d+\.\d+\.\d+)\.docx",
            r"(\d+\.\d+\.\d+)\.docx",
            r"(\d+\.\d+\.\d+)-\d+",
            r"(\d+\.\d+\.\d+)",
        ]:
            match = re.search(pattern, text)
            if match:
                return match.group(1)
    return ""


def parse_task_sections(tasks: str, qid: str) -> list[dict]:
    sections = []
    intro = tasks
    parts = re.split(r"(?=[（(]\d+[）)])", tasks)
    if len(parts) > 1:
        intro = parts[0].strip()
        for part in parts[1:]:
            head = re.match(r"[（(](\d+)[）)](.+)", part, re.DOTALL)
            if not head:
                continue
            num = head.group(1)
            body = head.group(2).strip()
            doc_match = re.search(r"题号为[“\"]([^”\"]+)[”\"]", body)
            doc_file_match = re.search(r"([\d.]+(?:-\d+)?)\.docx", body)
            jpg_match = re.search(r"命名为[“\"]([^”\"]+)[”\"]", body)
            deliverable = None
            if jpg_match:
                name = jpg_match.group(1)
                deliverable = name if "." in name else f"{name}.jpg"
            sections.append({
                "num": int(num),
                "text": body,
                "docSection": doc_match.group(1) if doc_match else None,
                "docFile": doc_file_match.group(0) if doc_file_match else None,
                "deliverable": deliverable,
            })
    return intro, sections


def find_docx_paths(qid: str, doc_name: str | None) -> list[Path]:
    names = set()
    if qid:
        names.add(f"{qid}.docx")
        names.add(f"素材{qid}.docx")
    if doc_name:
        names.add(doc_name)
    candidates = []
    for folder in [MATERIAL_DIR, ROOT / "qustions", QB_DIR / qid if qid else None]:
        if not folder or not folder.exists():
            continue
        for name in names:
            path = folder / name
            if path.exists():
                candidates.append(path)
        for name in names:
            for path in folder.glob(f"**/{name}"):
                candidates.append(path)
    return list(dict.fromkeys(candidates))


def parse_docx_template(path: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"source": str(path), "paragraphs": [], "fields": []}

    doc = Document(str(path))
    paragraphs = []
    fields = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        if re.search(r"_{2,}|【.*】|（\s*）|\(\s*\)|____", text):
            fields.append({
                "type": "line",
                "label": text,
                "placeholder": "",
                "prefill": text,
            })
        elif text.endswith("：") or text.endswith(":"):
            fields.append({
                "type": "textarea",
                "label": text,
                "placeholder": "请在此填写",
                "prefill": "",
            })

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
            for ci, cell_text in enumerate(cells):
                if not cell_text or re.search(r"_{2,}|待填|请填写", cell_text):
                    fields.append({
                        "type": "table_cell",
                        "label": f"表格·{cells[0] if cells else '单元格'}",
                        "placeholder": "",
                        "prefill": cell_text,
                    })
        paragraphs.append({"type": "table", "rows": rows})

    return {
        "source": path.name,
        "paragraphs": paragraphs,
        "fields": fields,
    }


def attach_code_rubric(code_items: list[dict], rubric: list[dict]) -> None:
    code_rubric = [r for r in rubric if "代码" in r["desc"]]
    blank_idx = 0
    for item in code_items:
        item_blanks = item.get("blanks", [])
        for blank in item_blanks:
            if blank_idx < len(code_rubric):
                rub = code_rubric[blank_idx]
                blank["rubricId"] = rub["id"]
                if not blank.get("score"):
                    blank["score"] = rub["score"]
            blank_idx += 1
        item["score"] = sum(b.get("score", 0) for b in item_blanks)
        if item_blanks:
            item["title"] = item_blanks[0].get("desc", item.get("title", ""))


def parse_ipynb_code_items(qid: str, rubric: list[dict] | None = None) -> list[dict]:
    ipynb_path = QB_DIR / qid / f"{qid}.ipynb"
    if not ipynb_path.exists():
        return []

    data = json.loads(ipynb_path.read_text(encoding="utf-8"))
    items = []

    for cell_idx, cell in enumerate(data.get("cells", [])):
        if cell.get("cell_type") != "code":
            continue
        source = "".join(cell.get("source", []))
        if not BLANK_PATTERN.search(source):
            continue

        lines = source.split("\n")
        pending_score = 0
        pending_desc = ""
        code_parts: list[str] = []
        blanks: list[dict] = []
        cell_title = ""

        for line in lines:
            if line.strip().startswith("#") and not BLANK_PATTERN.search(line):
                score_match = re.search(r"(\d+)分", line)
                if score_match:
                    pending_score = int(score_match.group(1))
                    pending_desc = re.sub(r"\d+分", "", line.strip("# ")).strip()
                    if not cell_title:
                        cell_title = pending_desc
                continue

            if BLANK_PATTERN.search(line):
                display_line = BLANK_PATTERN.sub("___", line)
                code_parts.append(display_line)
                count = len(BLANK_PATTERN.findall(line))
                base = pending_score // count if count else 0
                extra = pending_score % count if count else 0
                for bi in range(count):
                    blanks.append({
                        "blankIndex": len(blanks),
                        "score": base + (1 if bi < extra else 0),
                        "desc": pending_desc or display_line.strip(),
                        "answer": None,
                    })
                pending_score = 0
                pending_desc = ""
            else:
                code_parts.append(line)

        if not blanks:
            continue

        items.append({
            "id": f"code_{cell_idx}",
            "type": "code",
            "title": cell_title or f"代码块 {cell_idx + 1}",
            "code": "\n".join(code_parts).rstrip(),
            "cellIndex": cell_idx,
            "blanks": blanks,
            "score": sum(b["score"] for b in blanks),
            "answer": None,
        })

    if rubric:
        attach_code_rubric(items, rubric)
    return items


def is_doc_section(sec: dict) -> bool:
    if sec.get("docSection"):
        return True
    text = sec.get("text", "")
    if re.search(r"写在.+docx|docx答题卷|保存为\s*docx", text, re.I):
        return True
    if sec.get("docFile") and not sec.get("deliverable"):
        return True
    return False


def build_doc_items(qid: str, task_sections: list[dict], rubric: list[dict], intro: str) -> list[dict]:
    items = []
    doc_refs = re.findall(r"([\w.\-]+)\.docx", intro + "\n".join(s["text"] for s in task_sections))
    doc_refs = list(dict.fromkeys(doc_refs))

    for sec in task_sections:
        if not is_doc_section(sec):
            continue
        doc_file = sec.get("docFile") or (f"{qid}.docx" if qid else None)
        section_id = sec.get("docSection") or (f"{qid}-{sec['num']}" if qid else f"sec{sec['num']}")
        template = {"source": None, "paragraphs": [], "fields": []}

        for path in find_docx_paths(qid, doc_file):
            template = parse_docx_template(path)
            break

        if not template["fields"]:
            template["fields"] = [{
                "type": "textarea",
                "label": section_id,
                "placeholder": "请在此填写答案",
                "prefill": "",
            }]
            template["paragraphs"] = [sec["text"]]

        related = [r for r in rubric if section_id.replace("-", "") in r["desc"].replace(".", "") or section_id in r["desc"]]
        if not related:
            related = [r for r in rubric if str(sec["num"]) in r["desc"]]
        score = sum(r["score"] for r in related) if related else 0

        items.append({
            "id": f"doc_{section_id}",
            "type": "doc",
            "title": f"（{sec['num']}）简答题",
            "sectionId": section_id,
            "prompt": sec["text"],
            "docFile": doc_file,
            "template": template,
            "rubric": related,
            "score": score,
            "answer": None,
        })

    for doc_name in doc_refs:
        if any(i.get("docFile") == doc_name for i in items):
            continue
        template = {"source": None, "paragraphs": [], "fields": []}
        for path in find_docx_paths(qid, doc_name):
            template = parse_docx_template(path)
            break
        if not template["fields"]:
            template["fields"] = [{
                "type": "textarea",
                "label": "请补全以下内容",
                "placeholder": "请在此填写",
                "prefill": "",
            }]
        related = [r for r in rubric if doc_name.replace(".docx", "") in r["desc"]]
        items.append({
            "id": f"doc_{doc_name.replace('.docx', '')}",
            "type": "doc",
            "title": f"补全 {doc_name}",
            "sectionId": doc_name.replace(".docx", ""),
            "prompt": intro,
            "docFile": doc_name,
            "template": template,
            "rubric": related,
            "score": sum(r["score"] for r in related) if related else 0,
            "answer": None,
        })

    return items


def apply_manual_doc_template(qid: str, items: list[dict], rubric: list[dict]) -> list[dict]:
    """用手动模板替换自动解析的 doc 项（1.2.x / 2.1.x / 2.2.x 前 20 题）。"""
    if qid not in DOC_TEMPLATES:
        return items

    template = DOC_TEMPLATES[qid]
    doc_items = [i for i in items if i.get("type") == "doc"]
    other_items = [i for i in items if i.get("type") != "doc"]

    if template.get("skip"):
        return other_items

    related = []
    seen = set()
    if qid.startswith("3.2.") and not template.get("skip"):
        related = [r for r in rubric if "人机交互" in r["desc"]]
    elif qid.startswith("4.1."):
        related = [r for r in rubric if "回答内容" in r["desc"] or f"{qid}.docx" in r["desc"]]
    elif qid.startswith("4.2."):
        related = [r for r in rubric if "补充方案" in r["desc"] or "内容点" in r["desc"] or f"{qid}.docx" in r["desc"]]
    elif qid.startswith("2.1."):
        keywords = ["数据清洗规范", "数据标注规范", "特征工程规范"]
        related = [r for r in rubric if any(k in r["desc"] for k in keywords)]
    else:
        for item in doc_items:
            for rub in item.get("rubric", []):
                if rub["id"] not in seen:
                    related.append(rub)
                    seen.add(rub["id"])
        if not related:
            related = [r for r in rubric if f"{qid}.docx" in r["desc"] or qid in r["desc"]]

    title = "答题卷"
    if qid.startswith("1.2."):
        title = "简答题"
    elif qid.startswith("2.1."):
        title = "数据规范"
    elif qid.startswith("2.2."):
        title = "测试报告"
    elif qid.startswith("3.1."):
        title = "分析报告"
    elif qid.startswith("3.2."):
        title = "人机交互"
    elif qid.startswith("4.1."):
        title = "培训大纲"
    elif qid.startswith("4.2."):
        title = "指导方案"

    manual_doc = {
        "id": f"doc_{qid}",
        "type": "doc",
        "title": title,
        "sectionId": qid,
        "prompt": "",
        "docFile": f"{qid}.docx",
        "template": {
            "source": f"{qid}.docx",
            "layout": template["layout"],
            "sections": template["sections"],
        },
        "rubric": related,
        "score": sum(r["score"] for r in related),
        "answer": None,
        "docAnswer": get_index_answers().get(qid, []),
    }
    other_items.append(manual_doc)
    return other_items


def clean_tasks_text(tasks: str) -> str:
    text = tasks.strip()
    for marker in ("\n3.技能要求", "\n4.质量指标"):
        if marker in text:
            text = text.split(marker, 1)[0].strip()
    return text


def extract_between_tags(html: str, tag: str, class_name: str | None = None) -> str:
    if class_name:
        pattern = rf'<{tag}[^>]*class="{class_name}"[^>]*>(.*?)</{tag}>'
    else:
        pattern = rf'<{tag}[^>]*>(.*?)</{tag}>'
    match = re.search(pattern, html, re.DOTALL | re.IGNORECASE)
    return match.group(1).strip() if match else ""


def extract_html_sections(html: str) -> dict[str, str]:
    parts = re.split(r"<h3>", html, flags=re.IGNORECASE)
    sections: dict[str, str] = {}
    for part in parts[1:]:
        match = re.match(r"([^<]+)</h3>\s*(.*)", part, re.DOTALL | re.IGNORECASE)
        if match:
            sections[match.group(1).strip()] = match.group(2).strip()
    return sections


def find_question_html(qid: str) -> Path | None:
    matches = sorted(
        QUESTIONS_HTML_DIR.glob(f"{qid}*.html"),
        key=lambda p: (p.name != f"{qid}.html", p.name),
    )
    for path in matches:
        if path.name.startswith("_") or path.name == "index.html":
            continue
        return path
    return None


def html_to_plain_text(html_fragment: str) -> str:
    if not html_fragment:
        return ""

    text = html_fragment
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</p>\s*<p[^>]*>", "\n\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<p[^>]*>", "", text, flags=re.IGNORECASE)
    text = re.sub(r"</p>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = html_module.unescape(text)

    lines = []
    for raw_line in text.splitlines():
        line = raw_line.replace("\t", "").strip()
        if not line:
            continue
        if re.match(r"^[A-Za-z][A-Za-z0-9]*:\s", line):
            lines.append(f"• {line}")
        else:
            lines.append(line)
    return "\n".join(lines).strip()


def load_html_tasks(qid: str) -> str:
    path = find_question_html(qid)
    if not path or not path.exists():
        return ""

    html = path.read_text(encoding="utf-8")
    left_column = extract_between_tags(html, "div", "left-column")
    if not left_column:
        return ""

    sections = extract_html_sections(left_column)
    return html_to_plain_text(sections.get("工作任务", ""))


def build_task_description(intro: str, tasks: str, task_sections: list[dict]) -> str:
    cleaned = clean_tasks_text(tasks)
    intro_text = intro.strip()

    if intro_text and cleaned.startswith(intro_text):
        body = cleaned[len(intro_text):].strip()
        if body:
            return body

    if task_sections:
        use_fullwidth = "（1）" in cleaned or "（1）" in tasks
        open_paren = "（" if use_fullwidth else "("
        close_paren = "）" if use_fullwidth else ")"
        parts = []
        for sec in task_sections:
            text = sec.get("text", "").strip()
            if text:
                parts.append(f"{open_paren}{sec['num']}{close_paren}{text}")
        if parts:
            return "\n\n".join(parts).strip()

    return cleaned if cleaned != intro_text else ""


def attach_code_rubric_items(rubric: list[dict], code_items: list[dict]) -> None:
    code_rubric = [r for r in rubric if "代码" in r["desc"]]
    if not code_items or not code_rubric:
        return

    ci = 0
    for rub in code_rubric:
        if ci >= len(code_items):
            break
        code_items[ci]["rubric"] = [rub]
        if not code_items[ci].get("score"):
            code_items[ci]["score"] = rub["score"]
        ci += 1


def merge_question(name: str, sheet: str, rubric_block: str) -> dict:
    tasks = extract_field(sheet, "2.工作任务")
    time_match = re.search(r"考核时间[:：](.+)", sheet)
    time_text = time_match.group(1).strip() if time_match else ""

    rubric = parse_rubric(rubric_block)
    qid = infer_question_id(name, tasks, rubric)
    if not qid:
        raise ValueError(f"无法推断题号: {name}")

    intro, task_sections = parse_task_sections(tasks, qid)
    chapter = int(qid.split(".")[0]) if qid and qid[0].isdigit() else 0

    code_items = parse_ipynb_code_items(qid, rubric) if qid and (QB_DIR / qid).exists() else []
    doc_items = build_doc_items(qid, task_sections, rubric, intro)
    attach_code_rubric_items(rubric, code_items)
    task_content = load_html_tasks(qid)
    if not task_content.strip():
        intro, task_sections = parse_task_sections(tasks, qid)
        task_content = build_task_description(intro, tasks, task_sections) or clean_tasks_text(tasks)
    else:
        intro, task_sections = parse_task_sections(tasks, qid)

    items = []
    if task_content.strip():
        items.append({
            "id": "tasks",
            "type": "intro",
            "title": "工作任务",
            "content": task_content.strip(),
            "score": 0,
        })

    for sec in task_sections:
        if is_doc_section(sec):
            doc = next(
                (d for d in doc_items if d.get("sectionId") == sec.get("docSection") or d.get("sectionId") == f"{qid}-{sec['num']}"),
                None,
            )
            if doc and doc not in items:
                items.append(doc)

    for code in code_items:
        items.append(code)

    for doc in doc_items:
        if doc not in items:
            items.append(doc)

    items = apply_manual_doc_template(qid, items, rubric)

    total_score = sum(r["score"] for r in rubric) or sum(i.get("score", 0) for i in items)

    return {
        "id": qid,
        "code": qid,
        "name": name,
        "chapter": chapter,
        "chapterName": CHAPTER_NAMES.get(chapter, f"模块{chapter}"),
        "time": time_text,
        "tasks": tasks,
        "totalScore": total_score,
        "rubric": rubric,
        "items": items,
        "answer": None,
    }


def write_module(path: Path, data) -> None:
    path.write_text(
        f"module.exports = {json.dumps(data, ensure_ascii=False, indent=2)};\n",
        encoding="utf-8",
    )


THEORY_DIR = ROOT / "qustions" / "理论题"


def parse_judge_answers(text: str) -> dict[int, str]:
    """Parse 判断题答案.txt: ✓=正确(T), ✗=错误(F)."""
    answers: dict[int, str] = {}
    num = 0
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if re.match(r"^\d+[–-]\d+$", line):
            continue
        tokens = re.findall(r"[✓✗]", line)
        for token in tokens:
            num += 1
            answers[num] = "T" if token == "✓" else "F"
    return answers


def parse_judge_questions(text: str, answers: dict[int, str] | None = None) -> list[dict]:
    questions = []
    for match in re.finditer(r"（\s*）\s*(\d+)\.\s*(.+)", text):
        num = int(match.group(1))
        answer = None
        if answers and num in answers:
            answer = answers[num]
        questions.append({
            "id": f"j{num}",
            "num": num,
            "type": "judge",
            "text": match.group(2).strip(),
            "answer": answer,
        })
    return questions


def parse_choice_answers(text: str) -> dict[int, str]:
    """Parse 单选题答案.txt / 多选题答案.txt: '题号 答案' per line."""
    answers: dict[int, str] = {}
    for line in text.splitlines():
        line = line.strip()
        if not line or line.startswith("题号"):
            continue
        parts = line.split()
        if len(parts) < 2:
            continue
        if not parts[0].isdigit():
            continue
        num = int(parts[0])
        answer = parts[1].upper()
        answers[num] = answer
    return answers


def parse_choice_questions(
    text: str,
    qtype: str,
    prefix: str,
    answers: dict[int, str] | None = None,
) -> list[dict]:
    questions = []
    blocks = re.split(r"\n(?=\d+\.\s)", text)
    option_pattern = re.compile(r"^\(([A-E])\)\s*(.+)$")

    for block in blocks:
        block = block.strip()
        if not block:
            continue
        head = re.match(r"^(\d+)\.\s*(.+)$", block, re.DOTALL)
        if not head:
            continue
        num = int(head.group(1))
        lines = block.split("\n")
        stem = head.group(2).split("\n")[0].strip()
        options = {}
        for line in lines[1:]:
            line = line.strip()
            opt = option_pattern.match(line)
            if opt:
                options[opt.group(1)] = opt.group(2).strip()
        if not options:
            continue
        answer = None
        if answers and num in answers:
            raw = answers[num]
            if qtype == "multiple":
                answer = list(raw.upper())
            else:
                answer = raw.upper()
        questions.append({
            "id": f"{prefix}{num}",
            "num": num,
            "type": qtype,
            "text": stem,
            "options": options,
            "answer": answer,
        })
    return questions


def build_theory_data() -> dict:
    judge_answers = parse_judge_answers(read_text(THEORY_DIR / "判断题答案.txt"))
    judge = parse_judge_questions(read_text(THEORY_DIR / "判断题.md"), judge_answers)
    single_answers = parse_choice_answers(read_text(THEORY_DIR / "单选题答案.txt"))
    single = parse_choice_questions(
        read_text(THEORY_DIR / "单选题.md"), "single", "s", single_answers
    )
    multiple_answers = parse_choice_answers(read_text(THEORY_DIR / "多选题答案.txt"))
    multiple = parse_choice_questions(
        read_text(THEORY_DIR / "多选题.md"), "multiple", "m", multiple_answers
    )
    return {
        "judge": judge,
        "single": single,
        "multiple": multiple,
        "meta": {
            "title": "QuestMind",
            "subtitle": "人工智能训练师（三级）",
            "counts": {
                "judge": len(judge),
                "single": len(single),
                "multiple": len(multiple),
                "operation": 0,
                "total": len(judge) + len(single) + len(multiple),
            },
            "types": [
                {"key": "judge", "name": "判断题", "en": "True/False", "difficulty": "EASY", "icon": "fact_check"},
                {"key": "single", "name": "单选题", "en": "Single Choice", "difficulty": "MEDIUM", "icon": "radio_button_checked"},
                {"key": "multiple", "name": "多选题", "en": "Multiple Choice", "difficulty": "HARD", "icon": "checklist"},
                {"key": "operation", "name": "操作题", "en": "Hands-on Tasks", "difficulty": "EXPERT", "icon": "edit_note"},
            ],
        },
    }


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MATERIAL_DIR.mkdir(exist_ok=True)

    theory = build_theory_data()
    write_module(OUT_DIR / "judge.js", theory["judge"])
    write_module(OUT_DIR / "single.js", theory["single"])
    write_module(OUT_DIR / "multiple.js", theory["multiple"])

    text = read_text(OPS_MD)
    pairs = split_operation_md(text)
    questions = [normalize_deep(merge_question(name, sheet, rubric)) for name, sheet, rubric in pairs]
    questions.sort(key=lambda q: tuple(int(p) for p in q["id"].split(".")))

    meta = theory["meta"]
    meta["counts"]["operation"] = len(questions)
    meta["counts"]["total"] += len(questions)
    write_module(OUT_DIR / "meta.js", meta)
    write_module(OUT_DIR / "operation.js", questions)

    code_count = sum(1 for q in questions for i in q["items"] if i["type"] == "code")
    blank_count = sum(len(i.get("blanks", [])) for q in questions for i in q["items"] if i["type"] == "code")
    doc_count = sum(1 for q in questions for i in q["items"] if i["type"] == "doc")
    docx_found = sum(
        1 for q in questions for i in q["items"]
        if i["type"] == "doc" and i.get("template", {}).get("source")
    )

    print(f"判断题: {len(theory['judge'])}")
    print(f"单选题: {len(theory['single'])}")
    print(f"多选题: {len(theory['multiple'])}")
    print(f"操作题: {len(questions)}")
    print(f"  代码块: {code_count}, 填空: {blank_count}")
    print(f"  简答题(doc): {doc_count}, 已加载docx模板: {docx_found}")
    if doc_count and not docx_found:
        print(f"  提示: 将 docx 素材放入 {MATERIAL_DIR} 后重新运行可加载答题模板")
    print(f"输出: {OUT_DIR}")


if __name__ == "__main__":
    main()
